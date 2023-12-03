import csv
import datetime
import hashlib
import re
import shutil
import cpca

import fitz  # fitz就是pip install PyMuPDF
import cv2
import numpy as np
from docx import Document
from paddleocr import PaddleOCR, PPStructure, save_structure_res
# from paddleocr.ppstructure.predict_system import save_structure_res
from mysql_utils import *

# 数据库字段中英文对照字典
field_contrast = {"材料名称": "materialName",
                  "型号": "module",
                  "单位": "unit",
                  "数量": "number",
                  "含税价格": "taxPrice",
                  "不含税价格": "noTaxPrice",
                  "含增值税价格": "addTaxPrice",
                  "不含增值税价格": "noAddTaxPrice",
                  "签订时间": "signDate",
                  "甲方": "owner",
                  "乙方": "partyB",
                  "项目地址": "projectAddress",
                  "备注":"ps"
                  }


def write_data_to_csv(args):
    """
    传入记录文件名，处理文件路径，处理时间，文件处理结果
    输出：记录文件增加一行数据
    """
    file_path, name, process_time, file_status = args
    header = ["文件路径", "处理时间", "文件处理结果"]
    data = [name, process_time, file_status]
    # 如果文件存在，追加数据；否则，创建文件并写入数据
    mode = 'a' if os.path.exists(file_path) else 'w'
    with open(file_path, mode, newline='') as csv_file:
        csv_writer = csv.writer(csv_file)
        # 如果文件不存在，写入表头
        if mode == 'w':
            csv_writer.writerow(header)
        # 写入数据
        csv_writer.writerow(data)


def generate_md5(s):
    """
    字符串生成哈希值
    :param s:
    :return:
    """
    m = hashlib.md5()
    m.update(s.encode('utf-8'))
    return m.hexdigest()


def pyMuPDF_fitz(pdfPath, imagePath):
    """
    将pdf文件按页切分为图片
    """
    # images = convert_from_path(pdfPath, poppler_path="E:\\Release-23.11.0-0\\poppler-23.11.0\\Library\\bin")
    # images = convert_from_path(pdfPath, poppler_path=r"poppler-23.11.0\\Library\\bin")

    with fitz.open(pdfPath) as pdf:
        for pg in range(0, len(pdf)):
            page = pdf[pg]
            mat = fitz.Matrix(2, 2)
            pm = page.get_pixmap(matrix=mat, alpha=False)
            if pm.width > 2000 or pm.height > 2000:
                pm = page.get_pixmap(matrix=fitz.Matrix(1, 1), alpha=False)
            if not os.path.exists(imagePath):  # 判断存放图片的文件夹是否存在
                os.makedirs(imagePath)  # 若图片文件夹不存在就创建

            pm.save(imagePath + '/' + '%s.jpg' % pg)  # 将图片写入指定的文件夹内

def clean_num(s):
    s = str(s)
    num = re.findall("\d*\.?\d{0,2}", s)
    num = [s for s in num if s]
    if num:
        return num[0]
    return 0


def clean_data(table):
    """
    清洗价格表
    :param table:
    :return:
    """
    name_index = []
    title = []
    # 匹配指定名称的表头，获取索引，并重设表头
    # 判断是否是商品混凝土，需要区分处理
    if len(table)==0:
       return
    if re.findall("砼数量|基本单位", re.sub(",|\s", "", "".join(table[0]))):
        for s in table[0]:
            # 在匹配表头数据之前，替换掉换行和空格字符
            deal_s = re.sub(",|\s", "", s)
            # 匹配对应字段的位置，生成字段和字段索引两个列表
            if re.findall('强度等级', deal_s) and '型号' not in title:
                name_index.append(table[0].index(s))
                title.append('型号')
            elif re.findall("砼数量", deal_s) and '数量' not in title:
                name_index.append(table[0].index(s))
                title.append('数量')
            elif re.findall("基本单位", deal_s) and '含税价格' not in title:
                name_index.append(table[0].index(s))
                title.append('含税价格')
        # 获取指定列数据
        title = ['材料名称'] + title
        result_table = []
        for row in table[1:]:
            result_table.append(['商品混凝土'])
            for idx in name_index[:-1]:
                result_table[-1].append(row[idx])

    else:
        for s in table[0]:
            # 在匹配表头数据之前，替换掉换行和空格字符
            deal_s = re.sub(",|\s", "", s)
            # 匹配对应字段的位置，生成字段和字段索引两个列表
            if re.findall('名称', deal_s) and '材料名称' not in title:
                name_index.append(table[0].index(s))
                title.append('材料名称')
            elif re.findall("规格|型号|强度等级|洞口尺寸", deal_s) and '型号' not in title:
                name_index.append(table[0].index(s))
                title.append('型号')
            elif re.findall("单位", deal_s) and '单位' not in title:
                name_index.append(table[0].index(s))
                title.append('单位')
            elif re.findall("数量", deal_s) and '数量' not in title:
                name_index.append(table[0].index(s))
                title.append('数量')
            elif re.findall("单价", deal_s):
                if re.findall("不含增值税价格|不含增值税", deal_s) and '不含增值税价格' not in title:
                    name_index.append(table[0].index(s))
                    title.append('不含增值税价格')
                elif re.findall("含增值税价格|含增值税", deal_s) and '含增值税价格' not in title:
                    name_index.append(table[0].index(s))
                    title.append('含增值税价格')
                elif re.findall("不含税价格|不含税", deal_s) and '不含税价格' not in title:
                    name_index.append(table[0].index(s))
                    title.append('不含税价格')
                elif re.findall("含税价格|含税|单价", deal_s) and '含税价格' not in title:
                    name_index.append(table[0].index(s))
                    title.append('含税价格')
        # 获取指定列数据
        result_table = []
        for row in table[1:]:
            if re.findall("普通砼|泵送砼", "".join([str(s) for s in row])):
                continue
            try:
                # 判断合并表格，脏数据
                if len(set(row)) == 1:
                    continue
                # 生成行，并添加对应位置数据
                result_table.append([])
                for index, index_value in enumerate(name_index):
                    if title[index] == '含税价格' and re.findall("普通砼|泵送砼","".join([str(s) for s in table[1]])) and len(re.sub("[^\.\d]]","",str(row[index_value])))==0:
                        result_table[-1].append(row[index_value+1])
                    else:
                        result_table[-1].append(row[index_value])
            except Exception as e:
                print(f"error:{e}，{row}")
    # 返回清洗以后的表格数据
    return [title] + result_table


def d_date_clean(value):
    """
    一般日期处理
    """
    value = re.findall(r'\d+', value)
    if value:
        # 如果第一位不是年
        if len(value[0]) != 4:
            year = datetime.now().strftime('%Y')
            # 如果年是两位数
            if len(value[0]) == 2 and int(value[0]) > 12:
                value[0] = '20' + value[0]
            else:
                # 第一位是月份
                value.insert(0, year)
        if len(value) < 6:
            for i in range(6 - len(value)):
                value.append(0)
        # 全部转换成数字
        value = [int(i) for i in value]
        # 如果24点
        if value[3] == 24:
            value[3] = 23
            value[4] = 59
            value[5] = 59
        try:
            value = datetime(value[0], value[1], value[2], value[3], value[4], value[5])
        except:
            value = ''
        return value


def gettable_doc(output_path, file):
    # 打开.doc文件
    doc = Document(file)
    # 拼接所有段落文本数据
    text = "".join([para.text for para in doc.paragraphs])
    time_strs = re.findall("[原签订合同合约日期时间：：]{5,}.{8,15}?[日号]", re.sub("\s", "", text))
    if time_strs:
        result_time = datetime.strptime('1900-01-01', "%Y-%m-%d")
        for time_str in time_strs:
            time_str = re.sub("[^\d年月]", "", time_str)
            sign_date = d_date_clean(time_str)
            if sign_date and sign_date > result_time:
                result_time = sign_date
        if result_time == datetime.strptime('1900-01-01', "%Y-%m-%d"):
            result_time = ''
        print(result_time)
    else:
        result_time = ''
    # 遍历文档的表格
    for table in doc.tables:
        table = [[cell.text for cell in row.cells if cell] for row in table.rows]
        for s in table[0]:
            if re.findall("单价", re.sub(",|\s", "", str(s))):
                for i in range(len(table)):
                    # if [re.findall('合计|备注', s) for s in table[i] if '合计' in table[i]]:
                    if re.findall("合计|备注|总计|小计","".join([str(s) for s in table])):
                        table = table[:i]
                        table = clean_data(table)
                        df = pd.DataFrame(data=table[1:], columns=table[0])
                        if result_time:
                            df['签订时间'] = result_time
                        df = df.rename(columns=field_contrast)
                        df.to_excel(os.path.join(output_path, f'doc.xlsx'), index=False)
                        return df
    return pd.DataFrame()


def clean_output(output_path):
    for file in os.listdir(output_path):
        if ".xlsx" in file:
            file_name = os.path.join(output_path, file)
            df = pd.read_excel(file_name).fillna("")
            title = df.columns.tolist()
            for s in title:
                if re.findall("单价", re.sub(",|\s", "", str(s))):
                    table = [title] + df.values.tolist()
                    for i in range(len(table)):
                        if re.findall('合计|备注', "".join([str(s) for s in table[i]])):
                            table = table[:i]
                            break
                    if len(table) == 0:
                        return pd.DataFrame()
                    table = clean_data(table)
                    df = pd.DataFrame(data=table[1:], columns=table[0])
                    return df
        # os.remove(os.path.join(output_path, file))
    return pd.DataFrame()


def getTables(img_dir, output_path):
    result_time = datetime.strptime('1900-01-01', "%Y-%m-%d")
    df = pd.DataFrame()
    party_A = ''
    party_B = ''
    province = ''
    city = ''
    county = ''
    address = ''
    # 遍历图片
    for img in os.listdir(img_dir):
        img_path = os.path.join(img_dir, img)
        table_engine = PPStructure(show_log=False,use_gpu=True)#,type='STRUCTURE'
        # 含有中文路径的读取方式
        img = cv2.imdecode(np.fromfile(img_path, dtype=np.uint8), -1)
        # 灰度化
        image = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        # 二值化
        _, image = cv2.threshold(image, 0, 255, cv2.THRESH_OTSU)
        cv2.imwrite(img_path,image)

        # img = cv2.imread(img_path)
        # img = cv2.imdecode(np.fromfile(img_path, dtype=np.uint8), -1)
        result = table_engine(image)
        save_structure_res(result, "./", output_path)
        df = clean_output(output_path)
    for img in os.listdir(img_dir):
        img_path = os.path.join(img_dir, img)
        # 文字识别
        ocr = PaddleOCR(use_angle_cls=False, lang="ch")
        result = ocr.ocr(img_path, cls=False)
        res = result[0]  # 因为只有一张图片，所以结果只有1个，直接取出
        # 解析时间
        if res:
            # 拼接文本 ,解析签订时间
            page_text = "".join([r[-1][0] for r in res])
            # 根据规则找到可能匹配的文本
            time_strs = re.findall("[原签订合同合约日期时间：：]{5,}.{8,15}?[日号]", re.sub("\s", "", page_text))
            # 针对匹配到的时间进行清洗
            if time_strs:
                for time_str in time_strs:
                    # 去除了数字和时间之外的任何字符
                    time_str = re.sub("[^\d年月]", "", time_str)
                    # 清洗转换为事件类型
                    sign_date = d_date_clean(time_str)
                    # 找到最大的时间
                    if sign_date and sign_date > result_time:
                        result_time = sign_date
                        print(result_time)
            # 解析甲方乙方和项目地点
            for s in [r[-1][0] for r in res]:
                if re.findall("[甲需]方.{0,5}[：:](.*)", s) and party_A == "":
                    party_A = re.findall("[甲需]方.{0,5}[：:](.*)", s)[0]
                elif re.findall("[乙供]方.{0,5}[：:](.*)", s) and party_B == "":
                    party_B = re.findall("[乙供]方.{0,5}[：:](.*)", s)[0]
            # 解析项目地点
            page_text = "\n".join([r[-1][0] for r in res])
            if re.findall("项目\n地点.{0,5}[：:](.*)", page_text) and address == '':
                address = re.findall("项目\n?地点.{0,5}?[：:]\n?(.*)", page_text)[0]
                address_info = cpca.transform([address])
                if len(address_info) > 0:
                    province, city, county = address_info.values.tolist()[0][:3]

                # seg_list = list(jieba.cut(address, cut_all=True))
                # for s in seg_list:
                #     address_info = mysql_select_df(
                #         f"SELECT t0.ext_name AS 'sheng', t1.ext_name as 'shi', t2.ext_name as 'xian' FROM area T0 LEFT JOIN area T1 ON T0.id = T1.pid AND T1.deep = 1 LEFT JOIN area T2 ON T1.id = T2.pid AND T2.deep = 2 WHERE T0.deep = 0 and t2.ext_name like '%{s}%'")
                #     if len(address_info) > 0:
                #         province, city, county = address_info.values.tolist()[0]
                #         break
    # 判断是否存在材料单价表
    if len(df) != 0:
        # 判断是否解析到时间
        if result_time != datetime.strptime('1900-01-01', "%Y-%m-%d"):
            print(result_time)
            # 添加时间字段
            df['签订时间'] = result_time
        if party_A:
            df['甲方'] = party_A
        if party_B:
            df['乙方'] = party_B
        if address:
            df['项目地址'] = address
        if province:
            df['province'] = province
        if city:
            df['city'] = city
        if province:
            df['county'] = county
        # 字段重命名,替换为数据库对应字段
        df = df.rename(columns=field_contrast)
        # 存放到对应输出文件夹
        df.to_excel(os.path.join(output_path, f'pdf.xlsx'), index=False)
        return df
    return pd.DataFrame()


def gci(project_name, file):
    """
    传入项目名称和文件，根据文件后缀分流，进行处理
    :param project_name:
    :param file:
    :return:
    """
    # if file != r"\\192.168.180.180\13.平台-财务\平台公司共享\1.汇旺物资\合同扫描件\01安徽\10 2021年中央财政湿地保护资金项目-栖息地改造工程\板材-明鑫\10-11 汇旺-明鑫木业 栖息地项目-板材补充审批表+补充协议 180340元.pdf":
    #     return ''
    print('处理文件分割线'.center(100, '*'))
    print(file)
    file_path = file.replace("\\", "/")
    if file in pd.read_csv("file_log.csv",encoding='gbk')['文件路径'].values.tolist():
        print(f'文件已处理完成：{file}')
        return ""
    # 解析到的材料数据存放路径
    output_path = os.path.join('output', project_name, os.path.split(file)[-1].split(".")[0])
    # 创建存放路径在output文件夹中，下级目录项目名称，下级目录文件名称
    if not os.path.exists(output_path):
        os.makedirs(output_path)
    # 处理pdf文件
    if file.endswith(".pdf"):
        if os.path.exists(file.replace(".pdf", ".doc")) or os.path.exists(file.replace(".pdf", ".docx")):
            print("存在word文件，不进行pdf解析")
            return ""
        # image存放路径
        img_path = os.path.join('imgs', project_name, os.path.split(file)[-1].split(".")[0])
        # pdf生成图片
        pyMuPDF_fitz(file, img_path)
        # 清空临时输出文件
        if os.path.exists('test'):
            shutil.rmtree('test')
        # 解析pdf表格
        df = getTables(img_path, output_path)
    elif '.doc' in file:
        # doc文件和docx文件采用同一种处理方式
        df = gettable_doc(output_path, file)
    else:
        print(f"数据未处理：{file}")
        return ''

    """
    清洗数据入库
    """
    if len(df) != 0:
        # df = df.fillna('')
        df['projectName'] = project_name
        df['filePath'] = file_path
        df['createTime'] = datetime.now()
        df['updateTime'] = datetime.now()
        df['source'] = 1
        df = df.drop_duplicates()

        # 判断生成md5的字段的情况
        name_or_module = 'module'
        if 'materialName' in df.columns.tolist() and 'module' in df.columns.tolist():
            df = df[(df['materialName'] != '') | (df['module'] != '')]
        elif 'materialName' not in df.columns.tolist() and 'module' not in df.columns.tolist():
            print('文件解析材料数据结果无名称和规格！')
            write_data_to_csv(('file_log.csv', file, datetime.now(), '文件解析材料数据结果无名称和规格！'))
            return ''
        elif 'materialName' in df.columns.tolist():
            name_or_module = 'materialName'
        try:
            df['id'] = df.apply(lambda row: generate_md5(project_name + row['materialName'] + row['module']), axis=1)
        except Exception as e:
            print(f"error:容差性错误,不存在规格字段{e}")
            df['id'] = df.apply(lambda row: generate_md5(project_name + row[name_or_module]), axis=1)
        for col in df.columns:
            if col in ['number', 'taxPrice', 'noTaxPrice', 'addTaxPrice', 'noAddTaxPrice']:
                df[col] = df[col].apply(lambda cell: clean_num(cell))
        print(df)
        mysql_insert_data(df, 'material')
        print('文件有材料数据，已写入数据库！')
        write_data_to_csv(('file_log.csv', file, datetime.now(), '文件有材料数据，已写入数据库！'))
    else:
        print('文件无材料数据！')
        write_data_to_csv(('file_log.csv', file, datetime.now(), '文件无材料数据！'))


if __name__ == '__main__':
    # 项目
    pass
    # path = input("请输入合同文件夹: ")
    # for file1 in os.listdir(path):
    #     path1 = os.path.join(path, file1)
    #     project_name = re.findall("[^\d\s].*", file1)[0]
    #     print(project_name)
    #     # 合同类别
    #     for file2 in os.listdir(path1):
    #         path2 = os.path.join(path1, file2)
    #         if os.path.isdir(path2):
    #             # 合同文件
    #             for file3 in os.listdir(path2):
    #                 path3 = os.path.join(path2, file3)
    #                 if not re.findall("分包|施工|检测|机械", path3) and '合同' in path3 and (
    #                         path3.endswith(".pdf") or path3.endswith(".docx")):
    #                     gci(project_name, path3)
    #         elif (not re.findall("分包|施工|检测|机械", path2)) and '合同' in path2 and (
    #                 path2.endswith(".pdf") or path2.endswith(".docx")):
    #             gci(project_name, path2)
